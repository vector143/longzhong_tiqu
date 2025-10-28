import requests
import json
import time
import os
import re
from urllib.parse import quote, urljoin
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from qiniu import Auth, put_data
import csv
import datetime
from io import BytesIO
import base64
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
import queue
import random
from collections import defaultdict


# ==================== 增量更新日志记录器 ====================
class IncrementalUpdateLogger:
    """增量更新日志记录器"""

    def __init__(self, log_file="incremental_update.log"):
        self.log_file = log_file
        self.session_start_time = datetime.datetime.now()

    def log_session_start(self, keyword, days_back=None, hours_back=None):
        """记录会话开始"""
        time_range_desc = ""
        if hours_back is not None:
            time_range_desc = f"最近{hours_back}小时"
        elif days_back is not None:
            time_range_desc = f"最近{days_back}天"
        else:
            time_range_desc = "全部数据"

        log_entry = {
            'timestamp': self.session_start_time.strftime('%Y-%m-%d %H:%M:%S'),
            'type': 'SESSION_START',
            'keyword': keyword,
            'time_range': time_range_desc,
            'message': f"开始增量爬取: {keyword} ({time_range_desc})"
        }
        self._write_log(log_entry)
        print(f"📝 {log_entry['message']}")

    def log_articles_stats(self, articles):
        """记录文章统计信息"""
        if not articles:
            return

        # 分析发布时间范围 - 使用新的时间格式化方法
        publish_times = []
        for article in articles:
            publish_time = article.get('publishTime')
            if publish_time:
                # 使用新的时间格式化方法
                formatted_time = self._format_publish_time(publish_time)
                if formatted_time:
                    publish_times.append(formatted_time)

        if publish_times:
            min_time = min(publish_times)
            max_time = max(publish_times)
            time_range_msg = f"{min_time} 至 {max_time}"
        else:
            time_range_msg = "无发布时间信息"

        log_entry = {
            'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'type': 'ARTICLES_STATS',
            'total_articles': len(articles),
            'publish_time_range': time_range_msg,
            'html_files': sum(1 for article in articles if article['upload_status'].get('html')),
            'word_files': sum(1 for article in articles if article['upload_status'].get('word')),
            'message': f"爬取完成: 共{len(articles)}篇文章, 发布时间范围: {time_range_msg}"
        }
        self._write_log(log_entry)
        print(f"📊 {log_entry['message']}")

    def _format_publish_time(self, publish_time):
        """格式化发布时间 - 处理13位毫秒时间戳"""
        try:
            # 处理13位毫秒时间戳（如：1761441750966）
            if isinstance(publish_time, (int, float)) and publish_time > 1000000000000:
                timestamp_sec = publish_time / 1000
                dt = datetime.datetime.fromtimestamp(timestamp_sec)
                return dt.strftime('%Y-%m-%d %H:%M:%S')

            # 处理10位秒时间戳
            elif isinstance(publish_time, (int, float)) and publish_time > 1000000000:
                dt = datetime.datetime.fromtimestamp(publish_time)
                return dt.strftime('%Y-%m-%d %H:%M:%S')

            # 如果是字符串，尝试解析常见时间格式
            elif isinstance(publish_time, str):
                # 尝试常见的时间格式
                formats = [
                    '%Y-%m-%d %H:%M:%S',
                    '%Y/%m/%d %H:%M:%S',
                    '%Y-%m-%dT%H:%M:%S',
                    '%Y-%m-%d %H:%M',
                    '%Y-%m-%d'
                ]
                for fmt in formats:
                    try:
                        dt = datetime.datetime.strptime(publish_time, fmt)
                        return dt.strftime('%Y-%m-%d %H:%M:%S')
                    except ValueError:
                        continue

            # 其他无法识别的格式，返回原值
            return str(publish_time)

        except Exception as e:
            print(f"⚠️ 格式化发布时间失败: {publish_time}, 错误: {e}")
            return str(publish_time)

    def log_session_end(self, total_articles, elapsed_time):
        """记录会话结束"""
        log_entry = {
            'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'type': 'SESSION_END',
            'total_articles': total_articles,
            'elapsed_seconds': round(elapsed_time, 2),
            'message': f"会话结束: 处理{total_articles}篇文章, 耗时{elapsed_time:.2f}秒"
        }
        self._write_log(log_entry)
        print(f"✅ {log_entry['message']}")

    def log_error(self, error_message):
        """记录错误信息"""
        log_entry = {
            'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'type': 'ERROR',
            'message': error_message
        }
        self._write_log(log_entry)
        print(f"❌ 错误: {error_message}")

    def _write_log(self, log_entry):
        """写入日志文件"""
        try:
            with open(self.log_file, 'a', encoding='utf-8') as f:
                f.write(json.dumps(log_entry, ensure_ascii=False) + '\n')
        except Exception as e:
            print(f"⚠️ 写入日志失败: {e}")


# ==================== 上传任务类 ====================
class UploadTask:
    """上传任务封装"""

    def __init__(self, file_data, filename, file_type, article_info=None):
        self.file_data = file_data
        self.filename = filename
        self.file_type = file_type
        self.article_info = article_info or {}
        self.create_time = datetime.datetime.now()
        self.retry_count = 0

    def __str__(self):
        return f"UploadTask({self.filename}, {self.file_type}, size:{len(self.file_data)})"


# ==================== 异步内存七牛云上传器 ====================
class AsyncMemoryQiniuUploader:
    """异步内存七牛云上传器"""

    def __init__(self, access_key, secret_key, bucket_name, prefix="crawled_articles", max_upload_workers=3):
        self.auth = Auth(access_key, secret_key)
        self.bucket_name = bucket_name
        self.prefix = prefix

        # 上传队列和统计
        self.upload_queue = queue.Queue()
        self.uploaded_files = []
        self.failed_uploads = []
        self._lock = threading.Lock()
        self._stats_lock = threading.Lock()

        # 上传统计
        self.stats = {
            'total_submitted': 0,
            'total_success': 0,
            'total_failed': 0,
            'current_queue_size': 0,
            'upload_speed': 0
        }

        # 上传线程控制
        self.max_upload_workers = max_upload_workers
        self.upload_workers = []
        self.is_running = False

        print(f"✅ 异步上传器初始化: {max_upload_workers}个上传线程")

    def start_upload_workers(self):
        """启动上传工作线程"""
        if self.is_running:
            return

        self.is_running = True
        for i in range(self.max_upload_workers):
            worker = threading.Thread(
                target=self._upload_worker,
                name=f"UploadWorker-{i + 1}",
                daemon=True
            )
            worker.start()
            self.upload_workers.append(worker)

        print(f"🚀 启动 {self.max_upload_workers} 个上传线程")

    def stop_upload_workers(self):
        """停止上传工作线程"""
        self.is_running = False
        # 添加停止信号
        for _ in range(self.max_upload_workers):
            self.upload_queue.put(None)

        for worker in self.upload_workers:
            worker.join(timeout=5)

        self.upload_workers.clear()
        print("🛑 上传线程已停止")

    def submit_upload_task(self, file_data, filename, file_type, article_info=None):
        """提交上传任务到队列"""
        task = UploadTask(file_data, filename, file_type, article_info)

        with self._stats_lock:
            self.stats['total_submitted'] += 1
            self.stats['current_queue_size'] = self.upload_queue.qsize()

        self.upload_queue.put(task)
        print(f"📤 提交上传任务: {filename} (队列大小: {self.upload_queue.qsize()})")
        return True

    def _upload_worker(self):
        """上传工作线程"""
        thread_name = threading.current_thread().name

        while self.is_running:
            try:
                # 获取任务，超时1秒检查是否停止
                task = self.upload_queue.get(timeout=1)

                # 检查停止信号
                if task is None:
                    self.upload_queue.task_done()
                    break

                print(f"🧵 {thread_name} 处理上传: {task.filename}")

                # 执行上传
                success = self._execute_upload(task)

                # 更新统计
                with self._stats_lock:
                    if success:
                        self.stats['total_success'] += 1
                    else:
                        self.stats['total_failed'] += 1
                    self.stats['current_queue_size'] = self.upload_queue.qsize()

                self.upload_queue.task_done()

                # 短暂休息避免过于密集
                time.sleep(0.1)

            except queue.Empty:
                continue
            except Exception as e:
                print(f"❌ 上传工作线程异常: {e}")
                continue

    def _execute_upload(self, task):
        """执行实际上传操作"""
        try:
            # 构建安全的存储路径
            qiniu_key = self._build_safe_key(task.filename, task.file_type)

            print(f"📤 {threading.current_thread().name} 上传: {task.filename} -> {qiniu_key}")
            print(f"   数据大小: {len(task.file_data)} bytes, 重试次数: {task.retry_count}")

            # 生成上传token
            token = self.auth.upload_token(self.bucket_name, qiniu_key, 3600)

            # 直接从内存上传
            start_time = time.time()
            ret, info = put_data(token, qiniu_key, task.file_data)
            upload_time = time.time() - start_time

            print(f"📡 上传响应状态码: {info.status_code}, 耗时: {upload_time:.2f}s")

            if info.status_code == 200:
                print(f"✅ 上传成功: {qiniu_key}")
                with self._lock:
                    self.uploaded_files.append({
                        'filename': task.filename,
                        'qiniu_key': qiniu_key,
                        'file_type': task.file_type,
                        'upload_time': upload_time,
                        'article_info': task.article_info
                    })
                return True
            else:
                print(f"❌ 上传失败 {info.status_code}: {qiniu_key}")
                # 重试逻辑
                if task.retry_count < 2:  # 最多重试2次
                    task.retry_count += 1
                    print(f"🔄 重试上传 ({task.retry_count}/2): {task.filename}")
                    self.upload_queue.put(task)
                else:
                    with self._lock:
                        self.failed_uploads.append({
                            'filename': task.filename,
                            'error': f"状态码: {info.status_code}",
                            'retry_count': task.retry_count
                        })
                return False

        except Exception as e:
            print(f"💥 上传异常: {e}")
            # 重试逻辑
            if task.retry_count < 2:
                task.retry_count += 1
                print(f"🔄 重试上传 ({task.retry_count}/2): {task.filename}")
                self.upload_queue.put(task)
            else:
                with self._lock:
                    self.failed_uploads.append({
                        'filename': task.filename,
                        'error': str(e),
                        'retry_count': task.retry_count
                    })
            return False

    def _build_safe_key(self, filename, file_type):
        """构建安全的存储路径"""
        subdir = 'html' if file_type == 'html' else 'word'
        qiniu_key = f"{self.prefix}/{subdir}/{filename}".replace("\\", "/")
        return qiniu_key

    def wait_for_completion(self, timeout=300):
        """等待所有上传任务完成"""
        print(f"⏳ 等待上传队列清空... (当前队列: {self.upload_queue.qsize()})")

        start_time = time.time()
        while not self.upload_queue.empty():
            if time.time() - start_time > timeout:
                print(f"⚠️ 上传等待超时，剩余 {self.upload_queue.qsize()} 个任务")
                break

            remaining = self.upload_queue.qsize()
            if remaining > 0:
                print(f"📊 上传进度: 剩余 {remaining} 个任务")

            time.sleep(2)

        # 等待所有任务处理完成
        self.upload_queue.join()
        print("✅ 所有上传任务已完成")

    def get_upload_stats(self):
        """获取上传统计"""
        with self._stats_lock:
            stats = self.stats.copy()
            stats['uploaded_files_count'] = len(self.uploaded_files)
            stats['failed_uploads_count'] = len(self.failed_uploads)
            stats['queue_size'] = self.upload_queue.qsize()

        return stats

    def get_detailed_stats(self):
        """获取详细统计"""
        stats = self.get_upload_stats()

        # 按文件类型统计
        file_type_stats = defaultdict(int)
        for item in self.uploaded_files:
            file_type_stats[item['file_type']] += 1

        stats['file_type_breakdown'] = dict(file_type_stats)
        return stats


# ==================== 保留的源代码位置 1: UniversalNamingSystem ====================
# 【保留位置】从这里开始复制原有的 UniversalNamingSystem 类
class UniversalNamingSystem:
    """通用命名系统 - 线程安全版本"""

    def __init__(self, project_code="OIL"):
        self.project_code = project_code
        self.mapping_file = f"{project_code}_filename_mapping.csv"
        self.existing_mappings = self._load_existing_mappings()
        self._lock = threading.Lock()  # 添加线程锁

    def generate_universal_name(self, original_title, article_id, file_type='word'):
        """生成通用文件名：项目_日期_ID.扩展名"""
        # 提取日期
        date_str = self._extract_date(original_title)

        # 生成通用文件名
        ext = '.docx' if file_type == 'word' else '.html'
        universal_name = f"{self.project_code}_{date_str}_{article_id}{ext}"

        return universal_name

    def _extract_date(self, title):
        """从标题中提取日期"""
        patterns = [
            r'(\d{4})年(\d{1,2})月(\d{1,2})日',
            r'(\d{4})(\d{2})(\d{2})',
            r'(\d{4})-(\d{1,2})-(\d{1,2})',
        ]

        for pattern in patterns:
            match = re.search(pattern, title)
            if match:
                if '年' in pattern:
                    year, month, day = match.groups()
                else:
                    year, month, day = match.groups()
                return f"{year}{month.zfill(2)}{day.zfill(2)}"

        return datetime.datetime.now().strftime("%Y%m%d")

    def _load_existing_mappings(self):
        """加载已存在的映射关系"""
        mappings = set()
        if os.path.exists(self.mapping_file):
            try:
                df = pd.read_csv(self.mapping_file)
                mappings = set(df['new_filename'].tolist())
                print(f"📁 加载已有映射: {len(mappings)} 条记录")
            except Exception as e:
                print(f"⚠️ 加载映射文件失败: {e}")
        return mappings

    def save_mapping(self, original_title, new_filename, article_id, file_type, upload_status, local_saved=False, publish_time=None):
        """保存文件名映射到CSV（线程安全）"""
        with self._lock:  # 线程安全保护
            # 检查是否已存在
            if new_filename in self.existing_mappings:
                print(f"⏭️ 跳过已存在的映射: {new_filename}")
                return

            # 格式化发布时间
            formatted_publish_time = self._format_publish_time_for_csv(publish_time) if publish_time else "无发布时间"

            # 准备数据
            mapping_data = {
                'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'original_title': original_title,
                'new_filename': new_filename,
                'article_id': article_id,
                'file_type': file_type,
                'upload_status': upload_status,
                'local_saved': '是' if local_saved else '否',
                'publish_time': formatted_publish_time  # 新增列
            }

            # 写入CSV（追加模式）
            file_exists = os.path.exists(self.mapping_file)

            with open(self.mapping_file, 'a', newline='', encoding='utf-8-sig') as f:
                writer = csv.DictWriter(f, fieldnames=mapping_data.keys())

                if not file_exists:
                    writer.writeheader()  # 写入表头

                writer.writerow(mapping_data)

            # 更新已存在映射集合
            self.existing_mappings.add(new_filename)

            print(f"💾 保存映射: {new_filename}")

    def _format_publish_time_for_csv(self, publish_time):
        """为CSV格式化发布时间"""
        try:
            # 处理13位毫秒时间戳
            if isinstance(publish_time, (int, float)) and publish_time > 1000000000000:
                timestamp_sec = publish_time / 1000
                dt = datetime.datetime.fromtimestamp(timestamp_sec)
                return dt.strftime('%Y-%m-%d %H:%M:%S')

            # 处理10位秒时间戳
            elif isinstance(publish_time, (int, float)) and publish_time > 1000000000:
                dt = datetime.datetime.fromtimestamp(publish_time)
                return dt.strftime('%Y-%m-%d %H:%M:%S')

            # 如果是字符串，直接返回
            elif isinstance(publish_time, str):
                return publish_time

            # 其他格式返回字符串表示
            return str(publish_time)

        except Exception as e:
            print(f"⚠️ 格式化发布时间失败: {publish_time}, 错误: {e}")
            return str(publish_time)


# ==================== 异步格式转换管理器 ====================
class AsyncFormatConverter:
    def __init__(self, base_dir="articles", qiniu_uploader=None, naming_system=None, save_locally=False,
                 upload_to_qiniu=True):
        self.base_dir = base_dir
        self.qiniu_uploader = qiniu_uploader
        self.naming_system = naming_system
        self.save_locally = save_locally
        self.upload_to_qiniu = upload_to_qiniu  # 新增：七牛云上传开关
        self._lock = threading.Lock()

        if self.save_locally:
            self.html_dir = os.path.join(base_dir, "html")
            self.word_dir = os.path.join(base_dir, "word")
            self.images_dir = os.path.join(base_dir, "images")
            os.makedirs(self.html_dir, exist_ok=True)
            os.makedirs(self.word_dir, exist_ok=True)
            os.makedirs(self.images_dir, exist_ok=True)
            print(f"📁 本地目录已创建: {base_dir}")

        print(f"📤 七牛云上传: {'✅ 已启用' if self.upload_to_qiniu else '❌ 已禁用'}")

    # ... 其他方法保持不变 ...

    def save_as_html_async(self, html_content, title, article_id, publish_time = None):
        """异步保存为HTML文件"""
        try:
            # 生成安全文件名
            if self.naming_system:
                filename = self.naming_system.generate_universal_name(title, article_id, 'html')
            else:
                filename = f"{self.sanitize_filename(title)}_{article_id}.html"

            # 创建完整的HTML文档
            full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        table, th, td {{ border: 1px solid #ddd; }}
        th, td {{ padding: 8px; text-align: left; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {html_content}
</body>
</html>"""

            html_data = full_html.encode('utf-8')
            local_saved = False
            local_path = None

            # 保存到本地（如果启用）
            if self.save_locally:
                with self._lock:
                    filepath = os.path.join(self.html_dir, filename)
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(full_html)
                print(f"💾 HTML文件已保存到本地: {filename}")
                local_saved = True
                local_path = filepath
            else:
                print(f"🔧 生成HTML文件(内存): {filename}")

            print(f"✅ HTML文件生成完成: {len(html_data)} bytes")

            # 异步上传到七牛云（如果启用）
            upload_submitted = False
            if self.qiniu_uploader and self.upload_to_qiniu:  # 新增开关检查
                article_info = {
                    'title': title,
                    'article_id': article_id,
                    'file_type': 'html'
                }
                upload_submitted = self.qiniu_uploader.submit_upload_task(
                    html_data, filename, 'html', article_info
                )
                print(f"📤 HTML文件已提交上传: {filename}")
            elif self.upload_to_qiniu and not self.qiniu_uploader:
                print(f"⚠️ HTML文件跳过上传（无上传器）: {filename}")
            else:
                print(f"⏸️ HTML文件跳过上传（上传已禁用）: {filename}")

            # 保存映射关系
            if self.naming_system:
                self.naming_system.save_mapping(
                    title, filename, article_id, 'html',
                        'submitted' if upload_submitted else (
                            'disabled' if not self.upload_to_qiniu else 'no_uploader'),
                    local_saved,
                    publish_time=publish_time  # 新增参数
                )

            return {
                'filename': filename,
                'data_size': len(html_data),
                'upload_submitted': upload_submitted,
                'local_saved': local_saved,
                'local_path': local_path
            }

        except Exception as e:
            print(f"❌ HTML文件生成失败: {e}")
            if self.naming_system:
                filename = self.naming_system.generate_universal_name(title, article_id, 'html')
                upload_status = 'error' if self.upload_to_qiniu else 'disabled'
                self.naming_system.save_mapping(title, filename, article_id, 'html', upload_status, False, publish_time)
            return None

    def html_to_word_async(self, html_content, title, article_id, publish_time = None):
        """异步转换为Word文档"""
        try:
            # 生成安全文件名
            if self.naming_system:
                filename = self.naming_system.generate_universal_name(title, article_id, 'word')
            else:
                filename = f"{self.sanitize_filename(title)}_{article_id}.docx"

            # 在内存中创建Word文档
            doc = Document()
            doc.add_heading(title, level=1)

            soup = BeautifulSoup(html_content, 'html.parser')

            # 找到主要内容区域
            content_div = soup.find('div', class_='xq-content')
            if not content_div:
                content_div = soup

            print("🔍 开始按顺序处理内容...")

            # 处理内容元素
            elements_to_process = []
            for element in content_div.descendants:
                if not hasattr(element, 'name') or element.name is None:
                    continue

                element_name = element.name.lower()

                if element_name == 'p' and not element.find_parent('table'):
                    images_in_p = element.find_all('img')
                    if images_in_p:
                        elements_to_process.append(('image_paragraph', element, None))
                    else:
                        text = element.get_text(strip=True)
                        if text and text not in ['', '<br>']:
                            elements_to_process.append(('paragraph', element, text))
                elif element_name == 'table':
                    elements_to_process.append(('table', element, None))

            # 按文档顺序处理（去重）
            processed_elements = set()
            paragraph_count = 0
            table_count = 0
            div_text_count = 0
            image_count = 0

            for elem_type, element, text in elements_to_process:
                if id(element) not in processed_elements:
                    if elem_type == 'paragraph':
                        paragraph_count += 1
                        print(f"📝 段落 {paragraph_count}: '{text[:80]}...'")
                        doc.add_paragraph(text)
                    elif elem_type == 'table':
                        table_count += 1
                        print(f"📊 表格 {table_count}")
                        self._process_table_enhanced(element, doc)
                    # elif elem_type == 'div_text':
                    #     div_text_count += 1
                    #     print(f"📄 文本块 {div_text_count}: '{text[:100]}...'")
                    #     doc.add_paragraph(text)
                    elif elem_type == 'image_paragraph':  # NEW: 处理包含图片的段落
                        image_count += 1
                        print(f"🖼️  处理图片段落 {image_count}")
                        self._process_image_paragraph(element, doc)

                    processed_elements.add(id(element))

            if paragraph_count == 0 and table_count == 0 and image_count == 0:
                for element in content_div.descendants:
                    # 安全检查
                    if not hasattr(element, 'name') or element.name is None:
                        continue
                    element_name = element.name.lower()
                    if (element_name == 'div' and
                            element.get('id') == 'tableScroll' and
                            not element.find_parent('table')):
                        full_text = element.get_text(separator=' ', strip=True)
                        if full_text and len(full_text) > 10:
                            elements_to_process.append(('div_text', element, full_text))
                for elem_type, element, text in elements_to_process:
                    if elem_type == 'div_text':
                        div_text_count += 1
                        print(f"📄 文本块 {div_text_count}: '{text[:100]}...'")
                        doc.add_paragraph(text)

            print(
                f"✅ 处理完成: {paragraph_count} 个段落, {table_count} 个表格, {div_text_count} 个文本块, {image_count} 个图片段落")

            local_saved = False
            local_path = None

            # 保存到本地（如果启用）
            if self.save_locally:
                with self._lock:
                    filepath = os.path.join(self.word_dir, filename)
                    doc.save(filepath)
                print(f"💾 Word文件已保存到本地: {filename}")
                local_saved = True
                local_path = filepath
            else:
                print(f"🔧 生成Word文档(内存): {filename}")

            # 保存到内存用于上传
            memory_stream = BytesIO()
            doc.save(memory_stream)
            memory_stream.seek(0)
            docx_data = memory_stream.getvalue()
            memory_stream.close()

            print(f"✅ Word文档生成完成: {len(docx_data)} bytes")

            # 异步上传到七牛云（如果启用）
            upload_submitted = False
            if self.qiniu_uploader and self.upload_to_qiniu:  # 新增开关检查
                article_info = {
                    'title': title,
                    'article_id': article_id,
                    'file_type': 'word'
                }
                upload_submitted = self.qiniu_uploader.submit_upload_task(
                    docx_data, filename, 'word', article_info
                )
                print(f"📤 Word文件已提交上传: {filename}")
            elif self.upload_to_qiniu and not self.qiniu_uploader:
                print(f"⚠️ Word文件跳过上传（无上传器）: {filename}")
            else:
                print(f"⏸️ Word文件跳过上传（上传已禁用）: {filename}")

            # 保存映射关系
            if self.naming_system:
                self.naming_system.save_mapping(
                    title, filename, article_id, 'word',
                    'submitted' if upload_submitted else (
                        'disabled' if not self.upload_to_qiniu else 'no_uploader'),
                    local_saved,
                    publish_time=publish_time  # 新增参数
                )

            return {
                'filename': filename,
                'data_size': len(docx_data),
                'upload_submitted': upload_submitted,
                'local_saved': local_saved,
                'local_path': local_path
            }

        except Exception as e:
            print(f"❌ Word文档生成失败: {e}")
            if self.naming_system:
                filename = self.naming_system.generate_universal_name(title, article_id, 'word')
                upload_status = 'error' if self.upload_to_qiniu else 'disabled'
                self.naming_system.save_mapping(title, filename, article_id, 'word', upload_status, False, publish_time)
            return None

    # ==================== 保留的源代码位置 2: 文件处理方法 ====================
    # 【保留位置】从这里开始复制原有的文件处理方法
    def _process_table_enhanced(self, table_element, doc):
        """增强的表格处理方法"""
        try:
            rows = table_element.find_all('tr')
            if not rows:
                return

            max_cols = 0
            for row in rows:
                cols = 0
                for cell in row.find_all(['td', 'th']):
                    colspan = int(cell.get('colspan', 1))
                    cols += colspan
                max_cols = max(max_cols, cols)

            if max_cols == 0:
                return

            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'

            for row_idx, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                col_idx = 0

                for cell in cells:
                    while (row_idx < len(table.rows) and col_idx < max_cols and
                           table.cell(row_idx, col_idx).text.strip()):
                        col_idx += 1

                    if row_idx >= len(table.rows) or col_idx >= max_cols:
                        continue

                    colspan = int(cell.get('colspan', 1))
                    word_cell = table.cell(row_idx, col_idx)
                    self._process_table_cell_content(cell, word_cell, doc)
                    col_idx += colspan

        except Exception as e:
            print(f"❌ 处理表格失败: {e}")
            self._process_table_fallback(table_element, doc)

    def _process_table_cell_content(self, html_cell, word_cell, doc):
        """处理表格单元格内容"""
        try:
            word_cell.text = ""
            images = html_cell.find_all('img')
            if images:
                for img in images:
                    self._process_table_image(img, word_cell, doc)

            text_content = self._extract_cell_text(html_cell)
            if text_content:
                if images:
                    paragraph = word_cell.add_paragraph()
                    paragraph.add_run(text_content)
                else:
                    word_cell.text = text_content

        except Exception as e:
            print(f"❌ 处理单元格内容失败: {e}")
            text_content = html_cell.get_text(strip=True)
            if text_content:
                word_cell.text = text_content

    def _process_table_image(self, img_element, word_cell, doc):
        """处理表格中的图片"""
        try:
            img_src = img_element.get('src')
            if not img_src:
                return

            if img_src.startswith('//'):
                img_src = 'https:' + img_src
            elif img_src.startswith('/'):
                img_src = 'https://www.oilchem.net' + img_src

            img_response = requests.get(img_src, timeout=30)
            img_response.raise_for_status()

            image_stream = BytesIO(img_response.content)
            paragraph = word_cell.paragraphs[0] if word_cell.paragraphs else word_cell.add_paragraph()

            width = int(img_element.get('width', 200))
            height = int(img_element.get('height', 150))

            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(width / 96), height=Inches(height / 96))

        except Exception as e:
            print(f"❌ 处理表格图片失败: {e}")
            paragraph = word_cell.paragraphs[0] if word_cell.paragraphs else word_cell.add_paragraph()
            alt_text = img_element.get('alt', '图片')
            paragraph.add_run(f"[图片: {alt_text}]")

    def _extract_cell_text(self, cell):
        """提取单元格文本内容"""
        cell_copy = BeautifulSoup(str(cell), 'html.parser')
        for img in cell_copy.find_all('img'):
            img.decompose()
        text = cell_copy.get_text(strip=True, separator=' ')
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    def _process_table_fallback(self, table_element, doc):
        """表格处理的降级方案"""
        try:
            table_text = table_element.get_text(strip=True, separator=' | ')
            if table_text:
                doc.add_paragraph(f"[表格内容: {table_text}]")
        except Exception as e:
            print(f"❌ 表格降级处理也失败: {e}")

    def _process_image_paragraph(self, paragraph_element, doc):
        """处理包含图片的段落"""
        try:
            text_content = paragraph_element.get_text(strip=True)
            if text_content and text_content not in ['', '<br>']:
                doc.add_paragraph(text_content)

            images = paragraph_element.find_all('img')
            for img in images:
                self._process_standalone_image(img, doc)
        except Exception as e:
            print(f"❌ 处理图片段落失败: {e}")

    def _process_standalone_image(self, img_element, doc):
        """处理独立图片"""
        try:
            img_src = img_element.get('src')
            if not img_src:
                return

            if img_src.startswith('//'):
                img_src = 'https:' + img_src
            elif img_src.startswith('/'):
                img_src = 'https://www.oilchem.net' + img_src

            img_response = requests.get(img_src, timeout=30)
            img_response.raise_for_status()

            image_stream = BytesIO(img_response.content)
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1

            width = int(img_element.get('width', 400))
            height = int(img_element.get('height', 300))

            max_width = Inches(5.0)
            max_height = Inches(4.0)

            width_ratio = max_width / Inches(width / 96)
            height_ratio = max_height / Inches(height / 96)
            scale_ratio = min(width_ratio, height_ratio, 1.0)

            adjusted_width = Inches(width / 96) * scale_ratio
            adjusted_height = Inches(height / 96) * scale_ratio

            run = paragraph.add_run()
            run.add_picture(image_stream, width=adjusted_width, height=adjusted_height)

        except Exception as e:
            print(f"❌ 处理正文图片失败: {e}")
            paragraph = doc.add_paragraph()
            alt_text = img_element.get('alt', '图表')
            paragraph.add_run(f"[图片: {alt_text}]")
            paragraph.alignment = 1
    # 【保留位置结束】文件处理方法结束


# ==================== 保留的源代码位置 3: OilChemCookiesManager ====================
# 【保留位置】从这里开始复制原有的 OilChemCookiesManager 类
class OilChemCookiesManager:
    def __init__(self, cookies_file='cookies.json'):
        self.cookies_file = cookies_file
        self.session = requests.Session()
        self._lock = threading.Lock()
        self._setup_session()

    def _setup_session(self):
        """配置会话基础参数"""
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Referer': 'https://www.oilchem.net/'
        })

    def load_cookies(self):
        """从Cookie Editor导出的JSON文件加载cookies"""
        with self._lock:
            try:
                with open(self.cookies_file, 'r', encoding='utf-8') as f:
                    cookies_data = json.load(f)

                self.session.cookies.clear()
                cookies_loaded = 0
                for cookie in cookies_data:
                    try:
                        domain = cookie.get('domain', '')
                        if domain.startswith('.'):
                            domain = domain[1:]

                        if not domain or not cookie.get('name'):
                            continue

                        cookie_obj = requests.cookies.create_cookie(
                            name=cookie['name'],
                            value=cookie['value'],
                            domain=domain,
                            path=cookie.get('path', '/'),
                            secure=cookie.get('secure', False),
                            rest={'HttpOnly': cookie.get('httpOnly', False)}
                        )
                        self.session.cookies.set_cookie(cookie_obj)
                        cookies_loaded += 1

                    except Exception as e:
                        print(f"跳过无效cookie: {cookie.get('name', 'unknown')} - {e}")
                        continue

                print(f"✅ 已从 {self.cookies_file} 加载 {cookies_loaded} 个cookies")
                return True

            except FileNotFoundError:
                print(f"❌ Cookies文件不存在: {self.cookies_file}")
                return False
            except json.JSONDecodeError:
                print(f"❌ Cookies文件格式错误: {self.cookies_file}")
                return False
            except Exception as e:
                print(f"❌ 加载cookies失败: {e}")
                return False

    def validate_session(self, test_url="https://www.oilchem.net/25-1011-17-d80b2c132805eb10.html"):
        """验证会话是否有效"""
        try:
            print("🔍 验证cookies有效性...")
            response = self.session.get(test_url, timeout=10)
            response.encoding = 'utf-8'

            if "立即登录" in response.text:
                print("❌ Cookies已失效：检测到'立即登录'提示")
                return False
            else:
                print("✅ Cookies有效：可正常访问全文")
                return True

        except Exception as e:
            print(f"❌ 验证会话时出错: {e}")
            return False

    def get_export_instructions(self):
        """返回cookies导出指引"""
        return f"""
        🚀 Cookies 导出指引：

        1. 在Edge浏览器中安装 'Cookie Editor' 插件
        2. 访问 https://www.oilchem.net 并确保已登录
        3. 点击插件图标 → Export → JSON (默认格式)
        4. 保存为 '{self.cookies_file}' 文件
        5. 将文件放在代码同一目录下

        当前cookies文件: {self.cookies_file}
        """


# 【保留位置结束】OilChemCookiesManager 类结束

# ==================== 异步工作线程函数 ====================
def crawl_article_worker_async(article_data, session, output_formats, converter, delay=1):
    """异步版本的工作线程函数"""
    try:
        article = article_data['article']
        index = article_data['index']
        total = article_data['total']

        print(f"🧵 {threading.current_thread().name} 处理第 {index}/{total} 篇: {article['title']}")

        # 提取正文内容
        content_data = extract_article_content(article['url'], session=session)

        # 处理文件（异步上传）
        upload_status = {}
        local_files = {}

        if 'html' in output_formats and content_data['html_content'] and "需要登录" not in content_data['html_content']:
            html_result = converter.save_as_html_async(
                content_data['html_content'],
                article['title'],
                article.get('articleId', f'id_{index}'),
                publish_time = article.get('publishTime')  # 新增参数
            )
            if html_result:
                upload_status['html'] = html_result['upload_submitted']
                if html_result['local_saved']:
                    local_files['html'] = html_result['local_path']

        if 'word' in output_formats and content_data['html_content'] and "需要登录" not in content_data['html_content']:
            word_result = converter.html_to_word_async(
                content_data['html_content'],
                article['title'],
                article.get('articleId', f'id_{index}'),
                publish_time=article.get('publishTime')  # 新增参数
            )
            if word_result:
                upload_status['word'] = word_result['upload_submitted']
                if word_result['local_saved']:
                    local_files['word'] = word_result['local_path']

        # 整理数据
        result_data = {
            'articleId': article.get('articleId', ''),
            'title': article.get('title', ''),
            'publishTime': article.get('publishTime', ''),
            'url': article.get('url', ''),
            'columnName': article.get('columnName', ''),
            'content_preview': article.get('content', '')[:100] + '...',
            'upload_status': upload_status,
            'local_files': local_files,
            'images_count': len(content_data['images_data']),
            'has_tables': 'table' in content_data['html_content'].lower()
        }

        # 请求间延时
        time.sleep(delay + random.uniform(0.1, 0.5))

        return result_data

    except Exception as e:
        print(f"❌ 线程 {threading.current_thread().name} 处理文章失败: {e}")
        return None


# ==================== 异步多线程爬取主函数 ====================
def crawl_articles_async_multithread(keyword, pages=3, delay=2, cookies_manager=None,
                                     output_formats=None, days_back=None, hours_back=None,
                                     converter=None, max_crawl_workers=5, max_upload_workers=3):
    """异步多线程爬取文章 - 增强时间范围控制"""
    if output_formats is None:
        output_formats = ['html', 'word']

    all_articles = []
    all_article_data = []

    # 确定使用的session
    if cookies_manager and cookies_manager.session:
        session = cookies_manager.session
        print("🔐 使用cookies会话进行爬取")
    else:
        session = None
        print("⚠️  未使用cookies会话（可能无法访问全文）")

    # 第一步：获取所有文章列表
    print("📋 正在获取文章列表...")
    for page in range(1, pages + 1):
        print(f"获取第 {page} 页文章列表...")
        list_data = get_article_list(keyword, page_no=page, session=session,
                                     days_back=days_back, hours_back=hours_back)

        if not list_data:
            print(f"第 {page} 页没有数据")
            continue

        articles = list_data['response']['list']
        print(f"第 {page} 页找到 {len(articles)} 篇文章")

        # # 调试：查看第一篇文章的时间格式（只在第一页显示）
        # if articles and page == 1:
        #     sample_article = articles[0]
        #     publish_time = sample_article.get('publishTime')
        #     if publish_time:
        #         # 使用相同的时间格式化逻辑显示样例
        #         logger = IncrementalUpdateLogger()  # 临时创建用于格式化
        #         formatted_time = logger._format_publish_time(publish_time)
        #         print(f"🔍 发布时间样例: {publish_time} → {formatted_time}")
        #     else:
        #         print(f"🔍 第一篇文章无发布时间字段")

        for i, article in enumerate(articles):
            all_article_data.append({
                'article': article,
                'index': len(all_article_data) + 1,
                'total': len(articles) * pages
            })

    print(f"📊 总共获取到 {len(all_article_data)} 篇文章，开始异步多线程处理...")

    # 第二步：使用线程池处理文章
    with ThreadPoolExecutor(max_workers=max_crawl_workers) as executor:
        # 提交所有任务
        future_to_article = {
            executor.submit(
                crawl_article_worker_async,
                article_data,
                session,
                output_formats,
                converter,
                delay
            ): article_data for article_data in all_article_data
        }

        # 收集结果
        completed_count = 0
        for future in as_completed(future_to_article):
            article_data = future_to_article[future]
            try:
                result = future.result()
                if result:
                    all_articles.append(result)
                    completed_count += 1

                    # 实时显示上传队列状态
                    if converter.qiniu_uploader and converter.upload_to_qiniu:
                        stats = converter.qiniu_uploader.get_upload_stats()
                        print(
                            f"✅ 完成进度: {completed_count}/{len(all_article_data)} | 上传队列: {stats['queue_size']}")
                    else:
                        print(f"✅ 完成进度: {completed_count}/{len(all_article_data)}")

                else:
                    print(f"❌ 文章处理失败: {article_data['article']['title']}")
            except Exception as e:
                print(f"❌ 文章处理异常: {e}")

    return all_articles


# ==================== 保留的源代码位置 4: 网络请求函数 ====================
# 【保留位置】从这里开始复制原有的网络请求函数
def get_article_list(keyword, page_no=1, page_size=10, session=None, days_back=None, hours_back=None):
    """获取文章列表 - 增强时间范围控制"""
    url = "https://search.oilchem.net/article/search"
    params = {
        'keyword': keyword,
        'pageNo': page_no,
        'pageSize': page_size,
        'channelIds': '',
        'highlightFields': 'title,content'
    }

    # 时间范围计算
    if hours_back is not None:
        end_time = int(time.time() * 1000)
        start_time = int(end_time - (hours_back * 60 * 60 * 1000))  # 添加 int() 转换
        params['startTime'] = str(start_time)
        params['endTime'] = str(end_time)
        print(f"⏰ 时间范围: 最近 {hours_back} 小时 ({_format_timestamp(start_time)} 至 {_format_timestamp(end_time)})")
    elif days_back is not None:
        end_time = int(time.time() * 1000)
        start_time = int(end_time - (days_back * 24 * 60 * 60 * 1000))  # 同样修复
        params['startTime'] = str(start_time)
        params['endTime'] = str(end_time)
        print(f"⏰ 时间范围: 最近 {days_back} 天 ({_format_timestamp(start_time)} 至 {_format_timestamp(end_time)})")

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Referer': f'https://search.oilchem.net/article.html?keyword={quote(keyword)}'
    }

    try:
        if session:
            response = session.get(url, params=params, headers=headers)
        else:
            response = requests.get(url, params=params, headers=headers)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        print(f"获取文章列表失败: {e}")
        return None

def _format_timestamp(timestamp_ms):
    """格式化时间戳为可读字符串"""
    timestamp_sec = timestamp_ms / 1000
    return datetime.datetime.fromtimestamp(timestamp_sec).strftime('%Y-%m-%d %H:%M:%S')


def extract_article_content(article_url, session=None):
    """提取文章正文内容"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8'
    }

    try:
        if article_url.startswith('//'):
            article_url = 'https:' + article_url
        elif article_url.startswith('/'):
            article_url = 'https://www.oilchem.net' + article_url

        print(f"📖 正在提取: {article_url}")

        if session:
            response = session.get(article_url, headers=headers, timeout=10)
        else:
            response = requests.get(article_url, headers=headers, timeout=10)

        response.raise_for_status()
        response.encoding = 'utf-8'

        if "立即登录" in response.text:
            print("❌ 需要登录：检测到'立即登录'提示")
            return {
                'html_content': "需要登录才能访问全文",
                'images_data': [],
                'title': '需要登录'
            }

        soup = BeautifulSoup(response.text, 'html.parser')
        content_selectors = [
            '.article-content', '.content', '.main-content',
            '.detail-content', '.article-detail', '.news-content',
            '[class*="content"]', '[class*="article"]', '.text'
        ]

        content_elem = None
        for selector in content_selectors:
            content_elem = soup.select_one(selector)
            if content_elem:
                break

        if not content_elem:
            content_elem = soup.find('body')

        if content_elem:
            content_copy = BeautifulSoup(str(content_elem), 'html.parser')
            clean_selectors = [
                'script', 'style', 'nav', 'header', 'footer',
                '.ad', '.ads', '.advertisement', '.sidebar', '.side-bar',
                '.menu', '.navigation', '.nav', '.comment', '.comments',
                '.share', '.social', '.breadcrumb', '.breadcrumbs',
                '.pagination', '.page-nav', '.related', '.recommend',
                '.popup', '.modal', '.toolbar', '.tools',
                '.widget', '.plugin', '.banner', '.promotion',
                '[class*="ad"]', '[class*="banner"]',
                '[class*="menu"]', '[class*="nav"]', '[id*="ad"]'
            ]

            for clean_selector in clean_selectors:
                for elem in content_copy.select(clean_selector):
                    elem.decompose()

            title_elem = content_copy.find(['h1', 'h2']) or soup.find('title')
            title = title_elem.get_text(strip=True) if title_elem else "无标题"

            images_data = []
            for img_idx, img in enumerate(content_copy.find_all('img')):
                img_src = img.get('src')
                if img_src:
                    images_data.append({
                        'src': img_src,
                        'alt': img.get('alt', ''),
                        'index': img_idx
                    })

            return {
                'html_content': str(content_copy),
                'images_data': images_data,
                'title': title
            }

        return {
            'html_content': "正文提取失败",
            'images_data': [],
            'title': '提取失败'
        }

    except Exception as e:
        print(f"❌ 提取正文失败 {article_url}: {e}")
        return {
            'html_content': f"提取失败: {str(e)}",
            'images_data': [],
            'title': '提取失败'
        }


# 【保留位置结束】网络请求函数结束

# ==================== 保留的源代码位置 5: 保存结果函数 ====================
# 【保留位置】从这里开始复制原有的保存结果函数
def save_results(articles, keyword, qiniu_uploader=None, save_locally=False):
    """保存结果到文件，实现增量写入"""
    if not articles:
        print("没有数据可保存")
        return 0

    # 保存为JSON
    json_filename = f'{keyword}_articles_with_content.json'
    with open(json_filename, 'w', encoding='utf-8') as f:
        json.dump(articles, f, ensure_ascii=False, indent=2)
    print(f"✅ JSON数据已保存到: {json_filename}")

    # CSV文件名固定
    csv_filename = f'{keyword}_articles_summary.csv'

    # 准备CSV数据
    csv_data = []
    for article in articles:
        csv_data.append({
            '爬取时间': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            '文章ID': article.get('articleId', ''),
            '标题': article['title'],
            '发布时间': article['publishTime'],
            '栏目': article['columnName'],
            '正文预览': article['content_preview'],
            'HTML文件': '✅已上传' if article['upload_status'].get('html') else '❌失败',
            'Word文件': '✅已上传' if article['upload_status'].get('word') else '❌失败',
            '本地保存': '✅是' if article.get('local_files') else '❌否',
            '图片数量': article['images_count'],
            '包含表格': '是' if article['has_tables'] else '否',
            '链接': article['url']
        })

    # 增量写入CSV
    try:
        file_exists = os.path.exists(csv_filename)
        df = pd.DataFrame(csv_data)

        if file_exists:
            df.to_csv(csv_filename, mode='a', index=False, encoding='utf-8-sig', header=False)
            print(f"✅ CSV数据已追加到: {csv_filename}")
        else:
            df.to_csv(csv_filename, index=False, encoding='utf-8-sig')
            print(f"✅ CSV文件已创建: {csv_filename}")

    except Exception as e:
        print(f"❌ CSV文件保存失败: {e}")
        df.to_csv(csv_filename, index=False, encoding='utf-8-sig')
        print(f"✅ CSV文件已重新创建: {csv_filename}")

        # 输出统计信息
    html_success = sum(1 for article in articles if article['upload_status'].get('html'))
    word_success = sum(1 for article in articles if article['upload_status'].get('word'))
    local_saved = sum(1 for article in articles if article.get('local_files'))
    total_images = sum(article['images_count'] for article in articles)
    tables_count = sum(1 for article in articles if article['has_tables'])

    print(f"\n📊 文件处理统计:")
    print(f"   HTML文件: {html_success}篇成功")
    print(f"   Word文件: {word_success}篇成功")
    print(f"   本地保存: {local_saved}篇")
    print(f"   总图片数: {total_images}张")
    print(f"   包含表格: {tables_count}篇")

    # 七牛云上传统计（只有在有上传器时才显示）
    if qiniu_uploader:
        stats = qiniu_uploader.get_upload_stats()
        print(f"   ☁️ 七牛云上传: {stats['total_success']}个文件成功, {stats['total_failed']}个文件失败")
    else:
        print(f"   ☁️ 七牛云上传: 已禁用")

    return len(articles)


# ==================== 异步多线程主函数 ====================
def extract_from_keyword_async_multithread(keyword="原油", pages_to_crawl=3, delay_between_requests=2,
                                           use_cookies=True, output_formats=None, days_back=None, hours_back=None,
                                           # 新增 hours_back
                                           qiniu_config=None, save_locally=False, upload_to_qiniu=True,
                                           max_crawl_workers=5, max_upload_workers=3):
    """异步多线程主函数 - 增强时间范围控制和日志记录"""

    if output_formats is None:
        output_formats = ['html', 'word']

    # 初始化日志记录器
    logger = IncrementalUpdateLogger()

    # 记录会话开始
    logger.log_session_start(keyword, days_back=days_back, hours_back=hours_back)

    # 初始化异步上传器（只有在启用上传时才初始化）
    qiniu_uploader = None
    if qiniu_config and upload_to_qiniu:
        qiniu_uploader = AsyncMemoryQiniuUploader(
            qiniu_config['access_key'],
            qiniu_config['secret_key'],
            qiniu_config['bucket_name'],
            prefix=qiniu_config.get('prefix', 'crawled_articles'),
            max_upload_workers=max_upload_workers
        )
        qiniu_uploader.start_upload_workers()
        print("✅ 异步上传器初始化并启动成功")
    elif upload_to_qiniu and not qiniu_config:
        print("⚠️ 七牛云上传已启用，但未提供配置，跳过上传")
    else:
        print("⏸️ 七牛云上传已禁用")

    # 初始化其他组件
    naming_system = UniversalNamingSystem("OIL")

    converter = AsyncFormatConverter(
        base_dir="articles",
        qiniu_uploader=qiniu_uploader,
        naming_system=naming_system,
        save_locally=save_locally,
        upload_to_qiniu=upload_to_qiniu
    )

    cookies_manager = None
    if use_cookies:
        cookies_manager = OilChemCookiesManager('cookies_tang.json')
        if not cookies_manager.load_cookies():
            logger.log_error("Cookies加载失败")
            print(cookies_manager.get_export_instructions())
            return
        if not cookies_manager.validate_session():
            logger.log_error("Cookies验证失败")
            print(cookies_manager.get_export_instructions())
            return

    print(f"🚀 开始异步多线程爬取 '{keyword}' 相关文章...")

    # 显示时间范围信息
    if hours_back is not None:
        print(f"📅 时间范围: 最近 {hours_back} 小时")
    elif days_back is not None:
        print(f"📅 时间范围: 最近 {days_back} 天")
    else:
        print(f"📅 时间范围: 不限时间（所有数据）")

    print(f"📁 输出格式: {', '.join(output_formats)}")
    print(f"💾 本地保存: {'✅ 已启用' if save_locally else '❌ 已禁用'}")
    print(f"📤 七牛云上传: {'✅ 已启用' if upload_to_qiniu else '❌ 已禁用'}")
    print(f"🧵 爬取线程: {max_crawl_workers} 个")
    if upload_to_qiniu:
        print(f"📤 上传线程: {max_upload_workers} 个")
    print(f"📊 映射文件: {naming_system.mapping_file}")

    start_time = time.time()

    try:
        articles = crawl_articles_async_multithread(
            keyword=keyword,
            pages=pages_to_crawl,
            delay=delay_between_requests,
            cookies_manager=cookies_manager,
            output_formats=output_formats,
            days_back=days_back,
            hours_back=hours_back,  # 传递 hours_back
            converter=converter,
            max_crawl_workers=max_crawl_workers,
            max_upload_workers=max_upload_workers
        )

        # 等待所有上传任务完成（只有在启用上传时才等待）
        if qiniu_uploader and upload_to_qiniu:
            print("\n⏳ 爬取完成，等待上传队列清空...")
            qiniu_uploader.wait_for_completion()
        elif upload_to_qiniu:
            print("\n✅ 爬取完成（无上传任务）")
        else:
            print("\n✅ 爬取完成（上传已禁用）")

        end_time = time.time()
        elapsed_time = end_time - start_time

        if articles:
            count = save_results(articles, keyword, qiniu_uploader, save_locally)

            # 记录文章统计
            logger.log_articles_stats(articles)

            # 显示详细统计（只有在启用上传时才显示）
            if qiniu_uploader and upload_to_qiniu:
                stats = qiniu_uploader.get_detailed_stats()
                print(f"\n📊 异步上传统计详情:")
                print(f"   总提交任务: {stats['total_submitted']}")
                print(f"   成功上传: {stats['total_success']}")
                print(f"   失败上传: {stats['total_failed']}")
                print(f"   文件类型分布: {stats.get('file_type_breakdown', {})}")
            elif upload_to_qiniu:
                print(f"\n📊 上传状态: 已启用但无上传器配置")
            else:
                print(f"\n📊 上传状态: 已禁用")

            print(f"\n=== 异步多线程爬取完成 ===")
            print(f"共爬取 {count} 篇文章")
            print(f"总耗时: {elapsed_time:.2f} 秒")
            print(f"平均每篇文章: {elapsed_time / count:.2f} 秒")

            # 记录会话结束
            logger.log_session_end(count, elapsed_time)

        else:
            print("❌ 没有爬取到任何文章")
            logger.log_error("没有爬取到任何文章")

    except Exception as e:
        error_msg = f"爬取过程发生异常: {str(e)}"
        print(f"❌ {error_msg}")
        logger.log_error(error_msg)
    finally:
        # 确保上传线程被正确停止（只有在启用上传时才停止）
        if qiniu_uploader and upload_to_qiniu:
            qiniu_uploader.stop_upload_workers()


if __name__ == "__main__":
    # 七牛云配置
    qiniu_config = {
        'access_key': "SIQj8WGygyT7D8xN2EM0nrk0AA8huuO7e8xhymAA",
        'secret_key': "DNdy6U58Efglj4BRKooMZ25x71qa761_cvDoQ0Q_",
        'bucket_name': "yztrade",
        'prefix': "crawled_articles/multi_test"
    }

    # 使用异步多线程版本 - 支持小时级增量更新
    extract_from_keyword_async_multithread(
        keyword="原油",
        pages_to_crawl=1,
        delay_between_requests=1,
        use_cookies=True,
        output_formats=['html', 'word'],
        hours_back=None,  # 新增：获取最近1小时的数据
        qiniu_config=qiniu_config,
        save_locally=True,
        upload_to_qiniu=False,
        max_crawl_workers=3,
        max_upload_workers=3
    )