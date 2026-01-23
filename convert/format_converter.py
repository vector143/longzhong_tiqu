"""
异步格式转换器 - 支持HTML/Word/Markdown/JSON
"""

import json
import os
import re
import threading
from io import BytesIO
from typing import Any, Dict, Optional

from bs4 import BeautifulSoup
from docx import Document

from utils.time_utils import format_publish_time
from .html_utils import html_to_markdown, html_to_text_and_tables
from .word_processor import WordDocumentProcessor


class AsyncFormatConverter:
    """异步格式转换管理器"""

    def __init__(
        self,
        base_dir: str = "articles",
        qiniu_uploader=None,
        naming_system=None,
        save_locally: bool = False,
        upload_to_qiniu: bool = True,
    ):
        self.base_dir = base_dir
        self.qiniu_uploader = qiniu_uploader
        self.naming_system = naming_system
        self.save_locally = save_locally
        self.upload_to_qiniu = upload_to_qiniu
        self._lock = threading.Lock()

        if self.save_locally:
            self.markdown_dir = os.path.join(base_dir, "markdown")
            os.makedirs(self.markdown_dir, exist_ok=True)
            self.json_dir = os.path.join(base_dir, "json")
            os.makedirs(self.json_dir, exist_ok=True)
            self.html_dir = os.path.join(base_dir, "html")
            os.makedirs(self.html_dir, exist_ok=True)
            self.word_dir = os.path.join(base_dir, "word")
            os.makedirs(self.word_dir, exist_ok=True)
            print(f"📁 本地目录已创建: {base_dir}")

        print(f"📤 七牛云上传: {'✅ 已启用' if self.upload_to_qiniu else '❌ 已禁用'}")

    def save_as_markdown(
        self, html_content: str, title: str, article_id: str, publish_time=None
    ) -> Optional[Dict[str, Any]]:
        """保存为Markdown文件"""
        try:
            if "没有开通该新闻的页面权限" in html_content:
                print(f"⏭️ 跳过无权限文章: {title[:30]}...")
                return None

            filename = self._generate_filename(title, article_id, "html").replace(
                ".html", ".md"
            )
            markdown_content = html_to_markdown(html_content, title)

            if not markdown_content or len(markdown_content) < 50:
                print(f"⏭️ 内容过少，跳过: {title[:30]}...")
                return None

            local_saved, local_path = False, None
            if self.save_locally:
                with self._lock:
                    filepath = os.path.join(self.markdown_dir, filename)
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(markdown_content)
                print(f"💾 Markdown已保存: {filename}")
                local_saved, local_path = True, filepath

            self._save_mapping(
                title,
                filename,
                article_id,
                "markdown",
                "local_only",
                local_saved,
                publish_time,
            )
            return {
                "filename": filename,
                "data_size": len(markdown_content),
                "local_saved": local_saved,
                "local_path": local_path,
            }
        except Exception as e:
            print(f"❌ Markdown保存失败: {e}")
            return None

    def save_as_json(
        self,
        html_content: str,
        title: str,
        article_id: str,
        publish_time=None,
        url: str = None,
        column_name: str = None,
    ) -> Optional[Dict[str, Any]]:
        """保存为JSON文件"""
        try:
            if "没有开通该新闻的页面权限" in html_content:
                print(f"⏭️ 跳过无权限文章: {title[:30]}...")
                return None

            filename = self._generate_filename(title, article_id, "html").replace(
                ".html", ".json"
            )
            content_text, tables = html_to_text_and_tables(html_content)

            if not content_text or len(content_text) < 50:
                print(f"⏭️ 内容过少，跳过: {title[:30]}...")
                return None

            formatted_time = format_publish_time(publish_time) if publish_time else ""
            payload = {
                "articleId": article_id,
                "title": title,
                "publishTime": formatted_time,
                "url": url or "",
                "columnName": column_name or "",
                "source": "隆众资讯",
                "content": content_text,
                "tables": tables,
            }
            json_text = json.dumps(payload, ensure_ascii=False, indent=2)

            local_saved, local_path = False, None
            if self.save_locally:
                with self._lock:
                    filepath = os.path.join(self.json_dir, filename)
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(json_text)
                print(f"💾 JSON已保存: {filename}")
                local_saved, local_path = True, filepath

            self._save_mapping(
                title,
                filename,
                article_id,
                "json",
                "local_only",
                local_saved,
                publish_time,
            )
            return {
                "filename": filename,
                "data_size": len(json_text),
                "local_saved": local_saved,
                "local_path": local_path,
            }
        except Exception as e:
            print(f"❌ JSON保存失败: {e}")
            return None

    def save_as_html_async(
        self, html_content: str, title: str, article_id: str, publish_time=None
    ) -> Optional[Dict[str, Any]]:
        """异步保存为HTML文件"""
        try:
            filename = self._generate_filename(title, article_id, "html")
            full_html = self._create_html_document(title, html_content)
            html_data = full_html.encode("utf-8")

            local_saved, local_path = False, None
            if self.save_locally:
                with self._lock:
                    filepath = os.path.join(self.html_dir, filename)
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(full_html)
                print(f"💾 HTML文件已保存到本地: {filename}")
                local_saved, local_path = True, filepath

            upload_submitted = self._submit_upload(
                html_data, filename, "html", title, article_id
            )
            upload_status = (
                "submitted"
                if upload_submitted
                else ("disabled" if not self.upload_to_qiniu else "no_uploader")
            )
            self._save_mapping(
                title,
                filename,
                article_id,
                "html",
                upload_status,
                local_saved,
                publish_time,
            )

            return {
                "filename": filename,
                "data_size": len(html_data),
                "upload_submitted": upload_submitted,
                "local_saved": local_saved,
                "local_path": local_path,
            }
        except Exception as e:
            print(f"❌ HTML文件生成失败: {e}")
            return None

    def html_to_word_async(
        self, html_content: str, title: str, article_id: str, publish_time=None
    ) -> Optional[Dict[str, Any]]:
        """异步转换为Word文档"""
        try:
            filename = self._generate_filename(title, article_id, "word")
            doc = Document()
            doc.add_heading(title, level=1)

            soup = BeautifulSoup(html_content, "html.parser")
            content_div = soup.find("div", class_="xq-content") or soup

            processor = WordDocumentProcessor(doc)
            elements_to_process = self._collect_elements(content_div)
            processed_elements = set()
            counts = {"paragraph": 0, "table": 0, "image": 0}

            for elem_type, element, text in elements_to_process:
                if id(element) in processed_elements:
                    continue
                if elem_type == "paragraph":
                    counts["paragraph"] += 1
                    doc.add_paragraph(text)
                elif elem_type == "table":
                    counts["table"] += 1
                    processor.process_table_enhanced(element)
                elif elem_type == "image_paragraph":
                    counts["image"] += 1
                    processor.process_image_paragraph(element)
                processed_elements.add(id(element))

            print(
                f"✅ 处理完成: {counts['paragraph']} 段落, {counts['table']} 表格, {counts['image']} 图片"
            )

            local_saved, local_path = False, None
            if self.save_locally:
                with self._lock:
                    filepath = os.path.join(self.word_dir, filename)
                    doc.save(filepath)
                print(f"💾 Word文件已保存到本地: {filename}")
                local_saved, local_path = True, filepath

            memory_stream = BytesIO()
            doc.save(memory_stream)
            memory_stream.seek(0)
            docx_data = memory_stream.getvalue()
            memory_stream.close()

            upload_submitted = self._submit_upload(
                docx_data, filename, "word", title, article_id
            )
            upload_status = (
                "submitted"
                if upload_submitted
                else ("disabled" if not self.upload_to_qiniu else "no_uploader")
            )
            self._save_mapping(
                title,
                filename,
                article_id,
                "word",
                upload_status,
                local_saved,
                publish_time,
            )

            return {
                "filename": filename,
                "data_size": len(docx_data),
                "upload_submitted": upload_submitted,
                "local_saved": local_saved,
                "local_path": local_path,
            }
        except Exception as e:
            print(f"❌ Word文档生成失败: {e}")
            return None

    def _generate_filename(self, title: str, article_id: str, file_type: str) -> str:
        """生成文件名"""
        if self.naming_system:
            return self.naming_system.generate_universal_name(
                title, article_id, file_type
            )
        safe_title = re.sub(r'[\\/*?:"<>|]', "", title)[:50]
        ext = ".docx" if file_type == "word" else f".{file_type}"
        return f"{safe_title}_{article_id}{ext}"

    def _create_html_document(self, title: str, content: str) -> str:
        """创建完整HTML文档"""
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        table, th, td {{ border: 1px solid #ddd; }}
        th, td {{ padding: 8px; text-align: left; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {content}
</body>
</html>"""

    def _collect_elements(self, content_div):
        """收集需要处理的元素"""
        elements = []
        for element in content_div.descendants:
            if not hasattr(element, "name") or element.name is None:
                continue
            name = element.name.lower()
            if name == "p" and not element.find_parent("table"):
                if element.find_all("img"):
                    elements.append(("image_paragraph", element, None))
                else:
                    text = element.get_text(strip=True)
                    if text and text not in ["", "<br>"]:
                        elements.append(("paragraph", element, text))
            elif name == "table":
                elements.append(("table", element, None))
        return elements

    def _submit_upload(
        self, data: bytes, filename: str, file_type: str, title: str, article_id: str
    ) -> bool:
        """提交上传任务"""
        if self.qiniu_uploader and self.upload_to_qiniu:
            article_info = {
                "title": title,
                "article_id": article_id,
                "file_type": file_type,
            }
            return self.qiniu_uploader.submit_upload_task(
                data, filename, file_type, article_info
            )
        return False

    def _save_mapping(
        self,
        title: str,
        filename: str,
        article_id: str,
        file_type: str,
        upload_status: str,
        local_saved: bool,
        publish_time,
    ) -> None:
        """保存映射关系"""
        if self.naming_system:
            self.naming_system.save_mapping(
                title,
                filename,
                article_id,
                file_type,
                upload_status,
                local_saved,
                publish_time,
            )
