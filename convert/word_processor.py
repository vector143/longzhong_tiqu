"""
Word文档处理器 - 表格和图片处理
"""

import re
from io import BytesIO
from typing import Optional

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches


class WordDocumentProcessor:
    """Word文档内容处理器"""

    def __init__(self, doc: Document):
        self.doc = doc

    def process_table_enhanced(self, table_element) -> None:
        """增强的表格处理方法"""
        try:
            rows = table_element.find_all("tr")
            if not rows:
                return

            max_cols = 0
            for row in rows:
                cols = sum(
                    int(cell.get("colspan", 1)) for cell in row.find_all(["td", "th"])
                )
                max_cols = max(max_cols, cols)

            if max_cols == 0:
                return

            table = self.doc.add_table(rows=len(rows), cols=max_cols)
            table.style = "Table Grid"

            for row_idx, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                col_idx = 0
                for cell in cells:
                    while (
                        row_idx < len(table.rows)
                        and col_idx < max_cols
                        and table.cell(row_idx, col_idx).text.strip()
                    ):
                        col_idx += 1
                    if row_idx >= len(table.rows) or col_idx >= max_cols:
                        continue
                    colspan = int(cell.get("colspan", 1))
                    word_cell = table.cell(row_idx, col_idx)
                    self._process_table_cell_content(cell, word_cell)
                    col_idx += colspan
        except Exception as e:
            print(f"❌ 处理表格失败: {e}")
            self._process_table_fallback(table_element)

    def _process_table_cell_content(self, html_cell, word_cell) -> None:
        """处理表格单元格内容"""
        try:
            word_cell.text = ""
            images = html_cell.find_all("img")
            if images:
                for img in images:
                    self._process_table_image(img, word_cell)
            text_content = self._extract_cell_text(html_cell)
            if text_content:
                if images:
                    paragraph = word_cell.add_paragraph()
                    paragraph.add_run(text_content)
                else:
                    word_cell.text = text_content
        except Exception as e:
            print(f"❌ 处理单元格内容失败: {e}")
            word_cell.text = html_cell.get_text(strip=True)

    def _process_table_image(self, img_element, word_cell) -> None:
        """处理表格中的图片"""
        try:
            img_src = self._normalize_image_url(img_element.get("src"))
            if not img_src:
                return
            img_response = requests.get(img_src, timeout=30)
            img_response.raise_for_status()
            image_stream = BytesIO(img_response.content)
            paragraph = (
                word_cell.paragraphs[0]
                if word_cell.paragraphs
                else word_cell.add_paragraph()
            )
            width = int(img_element.get("width", 200))
            height = int(img_element.get("height", 150))
            run = paragraph.add_run()
            run.add_picture(
                image_stream, width=Inches(width / 96), height=Inches(height / 96)
            )
        except Exception as e:
            print(f"❌ 处理表格图片失败: {e}")
            paragraph = (
                word_cell.paragraphs[0]
                if word_cell.paragraphs
                else word_cell.add_paragraph()
            )
            paragraph.add_run(f"[图片: {img_element.get('alt', '图片')}]")

    def _extract_cell_text(self, cell) -> str:
        """提取单元格文本内容"""
        cell_copy = BeautifulSoup(str(cell), "html.parser")
        for img in cell_copy.find_all("img"):
            img.decompose()
        text = cell_copy.get_text(strip=True, separator=" ")
        return re.sub(r"\s+", " ", text).strip()

    def _process_table_fallback(self, table_element) -> None:
        """表格处理的降级方案"""
        try:
            table_text = table_element.get_text(strip=True, separator=" | ")
            if table_text:
                self.doc.add_paragraph(f"[表格内容: {table_text}]")
        except Exception as e:
            print(f"❌ 表格降级处理也失败: {e}")

    def process_image_paragraph(self, paragraph_element) -> None:
        """处理包含图片的段落"""
        try:
            text_content = paragraph_element.get_text(strip=True)
            if text_content and text_content not in ["", "<br>"]:
                self.doc.add_paragraph(text_content)
            for img in paragraph_element.find_all("img"):
                self._process_standalone_image(img)
        except Exception as e:
            print(f"❌ 处理图片段落失败: {e}")

    def _process_standalone_image(self, img_element) -> None:
        """处理独立图片"""
        try:
            img_src = self._normalize_image_url(img_element.get("src"))
            if not img_src:
                return
            img_response = requests.get(img_src, timeout=30)
            img_response.raise_for_status()
            image_stream = BytesIO(img_response.content)
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = 1
            width = int(img_element.get("width", 400))
            height = int(img_element.get("height", 300))
            max_width, max_height = Inches(5.0), Inches(4.0)
            width_ratio = max_width / Inches(width / 96)
            height_ratio = max_height / Inches(height / 96)
            scale_ratio = min(width_ratio, height_ratio, 1.0)
            adjusted_width = Inches(width / 96) * scale_ratio
            adjusted_height = Inches(height / 96) * scale_ratio
            run = paragraph.add_run()
            run.add_picture(image_stream, width=adjusted_width, height=adjusted_height)
        except Exception as e:
            print(f"❌ 处理正文图片失败: {e}")
            paragraph = self.doc.add_paragraph()
            paragraph.add_run(f"[图片: {img_element.get('alt', '图表')}]")
            paragraph.alignment = 1

    def _normalize_image_url(self, img_src: Optional[str]) -> Optional[str]:
        """规范化图片URL"""
        if not img_src:
            return None
        if img_src.startswith("//"):
            return "https:" + img_src
        elif img_src.startswith("/"):
            return "https://www.oilchem.net" + img_src
        return img_src
